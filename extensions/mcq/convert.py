import sys
import json

from utils import (
    success,
    info,
    warning,
    error
)

from utils import clean_json_string

def main() -> None:
    """
    Main function to convert an answer to a multiple choice question and save it to a .docx file.
    """
    info(message="Converting answer to multiple choice question...")

    # Check if the correct number of arguments is provided
    if len(sys.argv) != 2:
        error("Usage: python3 convert.py <answer>")
        sys.exit(1)

    answer: str = sys.argv[1]

    import os
    from docx import Document

    output_folder: str = "mcq/outputs"
    
    # Check if the output folder exists, if not, create it
    if not os.path.exists(output_folder):
        info(message=f"Folder '{output_folder}' does not exist. Creating the folder.")
        os.makedirs(output_folder)

    base_filename: str = "mcq_output"
    file_extension: str = ".docx"
    file_number: int = 1
    output_path: str = os.path.join(output_folder, f"{base_filename}_{file_number}{file_extension}")

    # Ensure the output file does not overwrite an existing file
    while os.path.exists(output_path):
        file_number += 1
        output_path = os.path.join(output_folder, f"{base_filename}_{file_number}{file_extension}")

    try:
        doc = Document()
        lines: list[str] = answer.split('\n')
        
        # Process each line and add appropriate content to the document
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
        
        # Save the document to the output path
        doc.save(output_path)
        success(message=f"Answer successfully written to {output_path}")
    except Exception as e:
        error(message=f"Failed to write answer to docx file: {str(e)}")

if __name__ == "__main__":
    main()