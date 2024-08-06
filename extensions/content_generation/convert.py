import sys
import subprocess
from utils import success, info, warning, error
from docx import Document
import os

def main() -> None:
    info("Converting answer to Word Document...")
    # if len(sys.argv) != 2:
    #     error("Usage: python3 convert.py <answer>")
    #     sys.exit(1)

    answer = sys.argv[1]
    sources = sys.argv[2:]

    output_folder = "extensions/content_generation/outputs"
    if not os.path.exists(output_folder):
        warning(f"Folder '{output_folder}' does not exist. Creating the folder.")
        os.makedirs(output_folder)

    base_filename = "content_generation_output"
    file_extension = ".docx"
    file_number = 1
    output_path = os.path.join(output_folder, f"{base_filename}_{file_number}{file_extension}")

    while os.path.exists(output_path):
        file_number += 1
        output_path = os.path.join(output_folder, f"{base_filename}_{file_number}{file_extension}")

    try:
        doc = Document()
        lines = answer.split('\n')
        for line in lines:
            if "Bot: #" in line:
                line = line.split('Bot: ')[-1]

            if line.lstrip().startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.lstrip().startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.lstrip().startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                # Ensure the line is XML compatible by removing any control characters
                line = ''.join([c for c in line if ord(c) > 31 or c in '\t\n\r\f\v'])
                doc.add_paragraph(line)

        try:
            if sources:
                doc.add_heading("References", level=1)
                print(sources)
                for source in sources:
                    # Ensure the source is XML compatible by removing any control characters
                    source = ''.join([c for c in source if ord(c) > 31 or c in '\t\n\r\f\v'])
                    doc.add_paragraph(source)
        except:
            pass

        doc.save(output_path)
        success(f"Answer successfully written to {output_path}")
        subprocess.run(["start", output_path], shell=True)
    except Exception as e:
        error(f"Failed to write answer to docx file: {str(e)}")

if __name__ == "__main__":
    main()